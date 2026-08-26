import sys
import os
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_borders(cell, top=True, bottom=True, left=True, right=True, color="CCCCCC", sz="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders_elm = OxmlElement('w:tcBorders')
    
    for side, active in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        edge = OxmlElement(f'w:{side}')
        edge.set(qn('w:val'), 'single' if active else 'none')
        edge.set(qn('w:sz'), sz)
        edge.set(qn('w:space'), '0')
        edge.set(qn('w:color'), color)
        borders_elm.append(edge)
        
    tcPr.append(borders_elm)

def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def convert_pdf_to_exact_docx(pdf_path, docx_path):
    doc = Document()
    doc_pdf = fitz.open(pdf_path)

    for page_idx, page in enumerate(doc_pdf):
        rect = page.rect
        page_width_in = rect.width / 72.0
        page_height_in = rect.height / 72.0

        # Create or adjust section margins to match exact page dimensions
        if page_idx == 0:
            section = doc.sections[0]
        else:
            section = doc.add_section()

        section.page_width = Inches(page_width_in)
        section.page_height = Inches(page_height_in)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

        # Detect Tables on the Page via PyMuPDF TableFinder
        tabs = page.find_tables()
        table_rects = [tab.bbox for tab in tabs]

        # Extract structured content blocks
        text_page = page.get_text("dict", flags=fitz.TEXT_DEHYPHENATE)
        blocks = text_page.get("blocks", [])

        # Process each block according to vertical Y-coordinate
        blocks = sorted(blocks, key=lambda b: (b.get("bbox", [0, 0, 0, 0])[1], b.get("bbox", [0, 0, 0, 0])[0]))

        handled_table_indices = set()

        for b in blocks:
            bbox = b.get("bbox", (0, 0, 0, 0))
            b_center_y = (bbox[1] + bbox[3]) / 2.0

            # 1. Check if block is part of a table
            in_table_idx = None
            for idx, tab in enumerate(tabs):
                tb_box = tab.bbox
                if tb_box[0] <= bbox[0] and tb_box[2] >= bbox[2] and tb_box[1] <= b_center_y <= tb_box[3]:
                    in_table_idx = idx
                    break

            if in_table_idx is not None:
                if in_table_idx not in handled_table_indices:
                    handled_table_indices.add(in_table_idx)
                    table_obj = tabs[in_table_idx]
                    df = table_obj.extract()

                    if df and len(df) > 0:
                        table = doc.add_table(rows=len(df), cols=len(df[0]))
                        table.alignment = WD_TABLE_ALIGNMENT.CENTER
                        table.autofit = True

                        for r_idx, row_vals in enumerate(df):
                            for c_idx, val in enumerate(row_vals):
                                cell = table.cell(r_idx, c_idx)
                                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                                set_cell_borders(cell, color="B0B0B0", sz="4")
                                set_cell_margins(cell, top=60, bottom=60, left=100, right=100)

                                p = cell.paragraphs[0]
                                p.paragraph_format.space_before = Pt(0)
                                p.paragraph_format.space_after = Pt(0)
                                p.paragraph_format.line_spacing = 1.0

                                if val:
                                    run = p.add_run(str(val).strip())
                                    run.font.name = "Calibri"
                                    run.font.size = Pt(9.5)
                                    if r_idx == 0:
                                        run.font.bold = True

                        doc.add_paragraph()
                continue

            # 2. Text Paragraph Reconstruction with Fonts, Sizes, Colors & Indents
            if b.get("type") == 0:  # Text block
                p = doc.add_paragraph()
                
                # Calculate Left Indentation from Page Margin
                left_pt = max(0, bbox[0] - 36)  # 36pt = 0.5 in left margin
                if left_pt > 15:
                    p.paragraph_format.left_indent = Pt(min(left_pt, 250))

                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.05

                for line in b.get("lines", []):
                    for span in line.get("spans", []):
                        text = span.get("text", "")
                        if not text:
                            continue

                        run = p.add_run(text)
                        
                        # Preserve exact font size
                        size = span.get("size", 10.0)
                        run.font.size = Pt(max(6.0, min(size, 48.0)))

                        # Preserve font styling
                        font_name = span.get("font", "").lower()
                        if "bold" in font_name or "black" in font_name or "heavy" in font_name:
                            run.font.bold = True
                        if "italic" in font_name or "oblique" in font_name:
                            run.font.italic = True

                        # Normalize Font Family
                        if "times" in font_name or "serif" in font_name:
                            run.font.name = "Times New Roman"
                        elif "arial" in font_name or "helvetica" in font_name:
                            run.font.name = "Arial"
                        elif "courier" in font_name or "mono" in font_name:
                            run.font.name = "Courier New"
                        else:
                            run.font.name = "Calibri"

                        # Preserve text color
                        color_int = span.get("color", 0)
                        r = (color_int >> 16) & 255
                        g = (color_int >> 8) & 255
                        b_val = color_int & 255
                        run.font.color.rgb = RGBColor(r, g, b_val)

            # 3. Image Block Handling
            elif b.get("type") == 1:  # Image block
                img_bytes = b.get("image")
                if img_bytes:
                    temp_img = os.path.join(os.path.dirname(docx_path), f"img_{page_idx}_{b.get('number', 0)}.png")
                    with open(temp_img, "wb") as f:
                        f.write(img_bytes)
                    try:
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        img_width_in = (bbox[2] - bbox[0]) / 72.0
                        p.add_run().add_picture(temp_img, width=Inches(min(img_width_in, 6.0)))
                    finally:
                        if os.path.exists(temp_img):
                            os.remove(temp_img)

    doc_pdf.close()
    doc.save(docx_path)
    return 0

if __name__ == '__main__':
    if len(sys.argv) < 3:
        sys.exit(1)
    
    input_pdf = sys.argv[1]
    output_docx = sys.argv[2]
    
    try:
        sys.exit(convert_pdf_to_exact_docx(input_pdf, output_docx))
    except Exception as e:
        sys.stderr.write(f"Conversion engine error: {str(e)}\n")
        sys.exit(1)